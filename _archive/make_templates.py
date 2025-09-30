from pathlib import Path
from docx import Document

ROOT = Path(__file__).resolve().parent.parent
TEMPL = ROOT / "templates"
TEMPL.mkdir(exist_ok=True, parents=True)

def a1_basic_template():
    """
    Generates a docxtpl-compliant .docx template.
    This version carefully constructs loops to be compatible with docxtpl's syntax
    and ensures Jinja2 tags are kept within single XML 'runs'.
    """
    doc = Document()

    # Helper to add a paragraph with a single run to avoid splitting tags
    def add_para_single_run(text, style=None):
        p = doc.add_paragraph(style=style)
        p.add_run(text)
        return p

    # --- Header ---
    add_para_single_run("{{ report_title }}", style='Heading 1')
    add_para_single_run("Company: {{ company_name }}")
    add_para_single_run("Date: {{ publication_date }}\tVersion: {{ version }}")

    # --- Simple Sections ---
    add_para_single_run("Executive Summary", style='Heading 2')
    add_para_single_run("{{ executive_summary }}")
    
    add_para_single_run("Introduction", style='Heading 2')
    add_para_single_run("{{ introduction }}")

    # --- Looping Section for 'sections' data ---
    add_para_single_run("Overview", style='Heading 2')
    add_para_single_run("{%p for section in sections %}")
    p = add_para_single_run("{{ section.title }}", style='Heading 3')
    add_para_single_run("{{ section.content }}")
    add_para_single_run("{%p endfor %}")

    # --- Table with Loop for 'sales_by_region' data ---
    add_para_single_run("Sales Details by Region", style='Heading 2')
    
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'

    # Helper to set cell text in a single run
    def set_cell_single_run(cell, text):
        # Clear existing content (a new cell has one empty paragraph)
        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        p.add_run(text)

    # Row 1: Header (static)
    hdr_cells = table.rows[0].cells
    set_cell_single_run(hdr_cells[0], 'Region')
    set_cell_single_run(hdr_cells[1], 'Revenue (k€)')
    set_cell_single_run(hdr_cells[2], 'Change vs Q2')

    # Row 2: Table row loop start tag. Must be in the first cell.
    loop_start_cells = table.rows[1].cells
    set_cell_single_run(loop_start_cells[0], '{%tr for row in sales_by_region %}')
    set_cell_single_run(loop_start_cells[1], '')
    set_cell_single_run(loop_start_cells[2], '')
    
    # Row 3: Template row with content placeholders
    content_cells = table.rows[2].cells
    set_cell_single_run(content_cells[0], '{{ row.region }}')
    set_cell_single_run(content_cells[1], '{{ row.revenue_k_eur }}')
    set_cell_single_run(content_cells[2], '{{ row.change_vs_q2 }}')
    
    # Row 4: Table row loop end tag. Must be in the first cell.
    loop_end_cells = table.rows[3].cells
    set_cell_single_run(loop_end_cells[0], '{%tr endfor %}')
    set_cell_single_run(loop_end_cells[1], '')
    set_cell_single_run(loop_end_cells[2], '')

    # --- Looping List for 'key_achievements' ---
    add_para_single_run("Key Achievements", style='Heading 2')
    add_para_single_run("{%p for item in key_achievements %}")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("{{ item }}")
    add_para_single_run("{%p endfor %}")
    
    # --- Final Section ---
    add_para_single_run("Conclusion", style='Heading 2')
    add_para_single_run("{{ conclusion }}")

    (TEMPL / "a1_basic_template.docx").unlink(missing_ok=True)
    doc.save(TEMPL / "a1_basic_template.docx")


def a2_richtext_template():
    doc = Document()
    doc.add_heading("RichText demo", level=1)
    doc.add_paragraph("Intro: {{ intro }}")
    doc.add_paragraph("{{ rich_paragraph }}")
    (TEMPL / "a2_richtext_template.docx").unlink(missing_ok=True)
    doc.save(TEMPL / "a2_richtext_template.docx")

def a3_images_template():
    doc = Document()
    doc.add_heading("Immagini con InlineImage", level=1)
    doc.add_paragraph("{{ product_image }}")
    doc.add_paragraph("Caption: {{ caption }}")
    (TEMPL / "a3_images_template.docx").unlink(missing_ok=True)
    doc.save(TEMPL / "a3_images_template.docx")

if __name__ == "__main__":
    a1_basic_template()
    a2_richtext_template()
    a3_images_template()
    print("Template generati in:", TEMPL)
