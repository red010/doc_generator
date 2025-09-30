#!/usr/bin/env python3
import os
import yaml
import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT, WD_UNDERLINE
from docx.enum.style import WD_STYLE_TYPE
import glob
import re

def load_config(config_path):
    """Carica la configurazione da un file YAML."""
    with open(config_path, 'r', encoding='utf-8') as file:
        return yaml.safe_load(file)

def hex_to_rgb(hex_color):
    """Converte un colore esadecimale in un oggetto RGBColor."""
    if hex_color.startswith('#'):
        hex_color = hex_color[1:]
    return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))

def apply_styles_to_document(document, styles_config):
    """Applica gli stili definiti nella configurazione al documento."""
    # Prima verifica che esistano tutti gli stili necessari
    required_styles = list(styles_config['styles'].keys())
    for style_name in required_styles:
        if style_name not in document.styles:
            print(f"Creazione dello stile '{style_name}' nel documento...")
            if style_name.startswith('Heading'):
                level = int(style_name.split()[-1])
                document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
                
    # Ora applica le proprietà agli stili
    for style_name, style_properties in styles_config['styles'].items():
        try:
            style = document.styles[style_name]
            font = style.font
            paragraph_format = style.paragraph_format
            
            # Impostazione del font
            if 'font' in style_properties:
                font.name = style_properties['font']
            
            # Impostazione della dimensione del font
            if 'size' in style_properties:
                font.size = Pt(style_properties['size'])
            
            # Impostazione del grassetto
            if 'bold' in style_properties:
                font.bold = style_properties['bold']
            
            # Impostazione della sottolineatura
            if 'underline' in style_properties:
                if style_properties['underline']:
                    font.underline = WD_UNDERLINE.SINGLE
                else:
                    font.underline = WD_UNDERLINE.NONE
            
            # Impostazione del colore
            if 'color' in style_properties:
                try:
                    font.color.rgb = hex_to_rgb(style_properties['color'])
                except ValueError as e:
                    print(f"Errore nel convertire il colore '{style_properties['color']}': {e}")
            
            # Impostazione dell'interlinea
            if 'line_spacing' in style_properties:
                line_spacing = style_properties['line_spacing']
                if isinstance(line_spacing, (int, float)):
                    paragraph_format.line_spacing = line_spacing
            
            # Impostazione dello spazio prima del paragrafo
            if 'space_before' in style_properties:
                paragraph_format.space_before = Pt(style_properties['space_before'])
            
            # Impostazione dello spazio dopo il paragrafo
            if 'space_after' in style_properties:
                paragraph_format.space_after = Pt(style_properties['space_after'])
            
            # Impostazione del rientro
            if 'indent' in style_properties:
                paragraph_format.first_line_indent = Inches(style_properties['indent'])
            
            # Impostazione del "keep with next"
            if 'keep_with_next' in style_properties:
                paragraph_format.keep_with_next = style_properties['keep_with_next']
            
            print(f"Stile '{style_name}' applicato con successo.")
            
        except KeyError as e:
            print(f"Attenzione: Lo stile '{style_name}' non è stato trovato nel documento. Errore: {e}")

def apply_inline_formatting(paragraph, text, element):
    """Applica la formattazione in linea (grassetto, corsivo, sottolineato)."""
    run = paragraph.add_run(text)
    
    # Gestione del grassetto
    if element.name == 'strong' or element.name == 'b':
        run.bold = True
    
    # Gestione del corsivo
    if element.name == 'em' or element.name == 'i':
        run.italic = True
    
    # Gestione del sottolineato
    if element.name == 'u':
        run.underline = WD_UNDERLINE.SINGLE
    
    return run

def process_inline_elements(paragraph, element):
    """Processa elementi inline come grassetto e corsivo all'interno di un elemento."""
    if element.name in ['strong', 'em', 'b', 'i', 'u']:
        return apply_inline_formatting(paragraph, element.get_text(), element)
    elif element.string:
        return paragraph.add_run(element.string)
    else:
        for child in element.children:
            if hasattr(child, 'name'):
                process_inline_elements(paragraph, child)
            elif child.string:
                paragraph.add_run(child.string)

def convert_html_to_docx(html_content, output_path, styles_config):
    """Converte contenuto HTML in un documento DOCX."""
    soup = BeautifulSoup(html_content, 'html.parser')
    document = Document()
    
    # Applicazione degli stili dal file di configurazione
    apply_styles_to_document(document, styles_config)
    
    # Gestione dei tag HTML
    for element in soup.find_all(True):
        if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            if element.parent.name not in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'li']:
                level = int(element.name[1])
                style_name = f"Heading {level}"
                
                # Verifica che lo stile esista
                if style_name not in document.styles:
                    print(f"Lo stile '{style_name}' non esiste nel documento. Utilizzo dello stile predefinito.")
                    style_name = f"Heading{level}"  # Prova con formato alternativo (senza spazio)
                    
                    if style_name not in document.styles:
                        print(f"Anche lo stile '{style_name}' non esiste. Utilizzo 'Normal'.")
                        style_name = "Normal"
                
                p = document.add_paragraph(style=style_name)
                print(f"Aggiunto paragrafo con stile '{style_name}'")
                
                # Processo gli elementi inline all'interno del titolo
                for child in element.children:
                    if hasattr(child, 'name'):
                        process_inline_elements(p, child)
                    elif child.string:
                        p.add_run(child.string)
        elif element.name == 'p':
            if element.parent.name not in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'li']:
                p = document.add_paragraph(style='Normal')
                
                # Processo gli elementi inline all'interno del paragrafo
                for child in element.children:
                    if hasattr(child, 'name'):
                        process_inline_elements(p, child)
                    elif child.string:
                        p.add_run(child.string)
        elif element.name == 'ul':
            for li in element.find_all('li', recursive=False):
                p = document.add_paragraph(style='List Bullet')
                
                # Processo gli elementi inline all'interno dell'elemento di lista
                for child in li.children:
                    if hasattr(child, 'name'):
                        process_inline_elements(p, child)
                    elif child.string:
                        p.add_run(child.string)
        elif element.name == 'ol':
            for li in element.find_all('li', recursive=False):
                p = document.add_paragraph(style='List Number')
                
                # Processo gli elementi inline all'interno dell'elemento di lista numerata
                for child in li.children:
                    if hasattr(child, 'name'):
                        process_inline_elements(p, child)
                    elif child.string:
                        p.add_run(child.string)
        elif element.name == 'blockquote':
            if element.parent.name not in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'li', 'blockquote']:
                for child in element.children:
                    if hasattr(child, 'name') and child.name == 'p':
                        p = document.add_paragraph(style='Quote')
                        
                        # Processo gli elementi inline all'interno della citazione
                        for grandchild in child.children:
                            if hasattr(grandchild, 'name'):
                                process_inline_elements(p, grandchild)
                            elif grandchild.string:
                                p.add_run(grandchild.string)
        elif element.name == 'pre':
            if element.parent.name not in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'li']:
                code = element.get_text()
                p = document.add_paragraph(code, style='No Spacing')
        elif element.name == 'table':
            if element.parent.name not in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'li', 'table']:
                rows = len(element.find_all('tr'))
                # Trova il numero massimo di celle in qualsiasi riga
                max_cells = 0
                for row in element.find_all('tr'):
                    cells = len(row.find_all(['td', 'th']))
                    max_cells = max(max_cells, cells)
                
                if rows > 0 and max_cells > 0:
                    table = document.add_table(rows=rows, cols=max_cells)
                    
                    # Riempire la tabella con i dati
                    for i, row in enumerate(element.find_all('tr')):
                        cells = row.find_all(['td', 'th'])
                        for j, cell in enumerate(cells):
                            if j < max_cells:  # Assicurarsi di non superare l'indice massimo
                                table_cell = table.cell(i, j)
                                p = table_cell.paragraphs[0]
                                
                                # Processo gli elementi inline all'interno della cella
                                for child in cell.children:
                                    if hasattr(child, 'name'):
                                        process_inline_elements(p, child)
                                    elif child.string:
                                        p.add_run(child.string)
    
    # Elenca tutti gli stili applicati nel documento
    print("Stili disponibili nel documento finale:")
    for style in document.styles:
        print(f" - {style.name}")
    
    # Salvataggio del documento
    document.save(output_path)

def convert_markdown_to_docx(md_path, output_path, styles_config):
    """Converte un file Markdown in DOCX."""
    with open(md_path, 'r', encoding='utf-8') as file:
        md_content = file.read()
    
    # Conversione Markdown in HTML
    html_content = markdown.markdown(
        md_content,
        extensions=['tables', 'fenced_code', 'nl2br', 'codehilite']
    )
    
    # Conversione HTML in DOCX
    convert_html_to_docx(html_content, output_path, styles_config)

def main():
    """Funzione principale."""
    # Assicurarsi che la directory di output esista
    os.makedirs('output_docx', exist_ok=True)
    
    # Caricamento della configurazione
    config = load_config('config.yaml')
    
    # Trovare tutti i file Markdown nella directory di input
    md_files = glob.glob('input_md/*.md')
    
    if not md_files:
        print("Nessun file Markdown trovato nella directory 'input_md'.")
        return
    
    # Conversione di ciascun file
    for md_file in md_files:
        base_name = os.path.basename(md_file)
        name_without_ext = os.path.splitext(base_name)[0]
        output_file = f"output_docx/{name_without_ext}.docx"
        
        print(f"Conversione di {md_file} in {output_file}...")
        convert_markdown_to_docx(md_file, output_file, config)
        print(f"Conversione completata: {output_file}")

if __name__ == "__main__":
    main() 