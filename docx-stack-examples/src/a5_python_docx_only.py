#!/usr/bin/env python
"""
Esempio: Generazione DOCX da Markdown senza template
Converte automaticamente docx_docxtpl_docxcompose.md in documento Word.
"""

import re
from pathlib import Path
from typing import Optional, List, Dict, Any
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Configuration
ROOT = Path(__file__).resolve().parent.parent
BUILD = ROOT / "build"
BUILD.mkdir(exist_ok=True)
SOURCE_MD = ROOT / "guide" / "docx_docxtpl_docxcompose.md"
OUT = BUILD / "out_a5_python_docx_only.docx"


class MarkdownToDocxConverter:
    """Converte contenuto Markdown in documento DOCX"""

    def __init__(self):
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        """Configura stili personalizzati per il documento"""
        # Stile per codice inline
        if 'Code' not in [s.name for s in self.doc.styles if s.type == WD_STYLE_TYPE.CHARACTER]:
            code_style = self.doc.styles.add_style('Code', WD_STYLE_TYPE.CHARACTER)
            code_style.font.name = 'Courier New'
            code_style.font.size = Pt(10)

        # Stile per blocchi di codice
        if 'CodeBlock' not in [s.name for s in self.doc.styles if s.type == WD_STYLE_TYPE.PARAGRAPH]:
            code_block_style = self.doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
            code_block_style.font.name = 'Courier New'
            code_block_style.font.size = Pt(9)
            code_block_style.paragraph_format.left_indent = Inches(0.25)
            code_block_style.paragraph_format.right_indent = Inches(0.25)

    def parse_markdown(self, markdown_content: str) -> List[Dict[str, Any]]:
        """Parsa il contenuto Markdown e restituisce struttura per DOCX"""
        lines = markdown_content.split('\n')
        elements = []
        i = 0

        while i < len(lines):
            line = lines[i].strip()

            # Frontmatter YAML
            if line == '---' and i == 0:
                frontmatter = self._parse_frontmatter(lines, i)
                elements.extend(frontmatter)
                i += frontmatter[-1].get('lines_consumed', 1) if frontmatter else 1
                continue

            # Titoli
            if line.startswith('#'):
                header_match = re.match(r'^#{1,6}\s+(.+)$', line)
                if header_match:
                    level = len(line.split()[0])
                    title = header_match.group(1)
                    elements.append({
                        'type': 'heading',
                        'level': min(level, 9),  # DOCX supporta fino al livello 9
                        'text': title
                    })

            # Blocchi di codice (fenced code blocks)
            elif line.startswith('```'):
                code_block = self._parse_code_block(lines, i)
                if code_block:
                    elements.append(code_block)
                    i += code_block.get('lines_consumed', 1)
                    continue

            # Liste non ordinate
            elif line.startswith(('* ', '- ')):
                list_items = self._parse_list(lines, i, ['* ', '- '])
                elements.extend(list_items)
                i += len(list_items)
                continue

            # Liste ordinate
            elif re.match(r'^\d+\.\s+', line):
                list_items = self._parse_list(lines, i, [r'^\d+\.\s+'])
                elements.extend(list_items)
                i += len(list_items)
                continue

            # Linee vuote (ignoriamo)
            elif not line:
                pass

            # Paragrafi normali
            else:
                # Accumula righe consecutive in paragrafi
                paragraph_lines = []
                while i < len(lines) and lines[i].strip() and not lines[i].startswith('#') and not lines[i].startswith('```'):
                    paragraph_lines.append(lines[i])
                    i += 1
                    if i < len(lines) and lines[i].strip() == '':
                        break

                if paragraph_lines:
                    text = ' '.join(line.strip() for line in paragraph_lines if line.strip())
                    if text:
                        elements.append({
                            'type': 'paragraph',
                            'text': self._parse_inline_markdown(text)
                        })
                else:
                    i += 1
                    continue

            i += 1

        return elements

    def _parse_frontmatter(self, lines: List[str], start_idx: int) -> List[Dict[str, Any]]:
        """Parsa il frontmatter YAML"""
        elements = []
        i = start_idx + 1

        # Salta fino alla fine del frontmatter
        while i < len(lines) and lines[i].strip() != '---':
            line = lines[i].strip()
            if ':' in line and not line.startswith('#'):
                key, value = line.split(':', 1)
                key = key.strip()
                value = value.strip().strip('"').strip("'")

                if key == 'title':
                    elements.append({
                        'type': 'heading',
                        'level': 0,
                        'text': value
                    })
                elif key == 'description':
                    elements.append({
                        'type': 'paragraph',
                        'text': value,
                        'style': 'italic'
                    })
                elif key == 'authors':
                    if value and value != 'null':
                        elements.append({
                            'type': 'paragraph',
                            'text': f"Autori: {value}"
                        })
                elif key == 'last_updated':
                    elements.append({
                        'type': 'paragraph',
                        'text': f"Ultimo aggiornamento: {value}",
                        'style': 'caption'
                    })
            i += 1

        return elements

    def _parse_code_block(self, lines: List[str], start_idx: int) -> Optional[Dict[str, Any]]:
        """Parsa un blocco di codice delimitato da ```"""
        if not lines[start_idx].startswith('```'):
            return None

        i = start_idx + 1
        code_lines = []

        while i < len(lines) and not lines[i].startswith('```'):
            code_lines.append(lines[i])
            i += 1

        if code_lines:
            return {
                'type': 'code_block',
                'code': '\n'.join(code_lines),
                'lines_consumed': i - start_idx + 1
            }

        return None

    def _parse_list(self, lines: List[str], start_idx: int, markers: List[str]) -> List[Dict[str, Any]]:
        """Parsa liste ordinate e non ordinate"""
        items = []
        i = start_idx

        while i < len(lines):
            line = lines[i].strip()
            is_list_item = False

            for marker in markers:
                if isinstance(marker, str):
                    if line.startswith(marker):
                        text = line[len(marker):].strip()
                        is_list_item = True
                        break
                else:  # regex
                    match = re.match(marker, line)
                    if match:
                        text = line[match.end():].strip()
                        is_list_item = True
                        break

            if is_list_item:
                items.append({
                    'type': 'list_item',
                    'text': self._parse_inline_markdown(text),
                    'ordered': isinstance(markers[0], str) and markers[0][0].isdigit()
                })
                i += 1
            else:
                break

        return items

    def _parse_inline_markdown(self, text: str) -> str:
        """Parsa elementi Markdown inline (grassetto, corsivo, codice)"""
        # Rimuovi link markdown [text](url) -> text
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)

        # Rimuovi riferimenti [1]: -> nulla
        text = re.sub(r'\[\d+\]:\s*.*?(?=\n|$)', '', text, flags=re.MULTILINE)

        # Rimuovi backticks per codice inline (li gestiamo diversamente)
        text = re.sub(r'`([^`]+)`', r'\1', text)

        return text.strip()

    def convert_to_docx(self, elements: List[Dict[str, Any]]) -> Document:
        """Converte gli elementi parsati in documento DOCX"""
        current_list_items = []

        for element in elements:
            if element['type'] == 'heading':
                # Gestisci eventuali list items pendenti
                if current_list_items:
                    self._add_list(current_list_items)
                    current_list_items = []

                level = element['level']
                if level == 0:
                    # Titolo principale
                    title = self.doc.add_heading(element['text'], 0)
                    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    self.doc.add_heading(element['text'], level)

            elif element['type'] == 'paragraph':
                # Gestisci eventuali list items pendenti
                if current_list_items:
                    self._add_list(current_list_items)
                    current_list_items = []

                para = self.doc.add_paragraph(element['text'])
                if element.get('style') == 'italic':
                    for run in para.runs:
                        run.italic = True
                elif element.get('style') == 'caption':
                    para.style = 'Caption'

            elif element['type'] == 'list_item':
                current_list_items.append(element)

            elif element['type'] == 'code_block':
                # Gestisci eventuali list items pendenti
                if current_list_items:
                    self._add_list(current_list_items)
                    current_list_items = []

                # Aggiungi blocco di codice
                code_para = self.doc.add_paragraph(element['code'], style='CodeBlock')

        # Gestisci eventuali list items rimanenti
        if current_list_items:
            self._add_list(current_list_items)

        return self.doc

    def _add_list(self, items: List[Dict[str, Any]]):
        """Aggiunge una lista al documento"""
        for item in items:
            self.doc.add_paragraph(item['text'], style='List Bullet')


def convert_markdown_to_docx(markdown_path: Path, output_path: Optional[Path] = None) -> Path:
    """
    Converte un file Markdown in documento DOCX senza usare template.
    """
    if output_path is None:
        output_path = OUT

    print(f"📖 Lettura file Markdown: {markdown_path}")
    with open(markdown_path, 'r', encoding='utf-8') as f:
        markdown_content = f.read()

    print("🔄 Conversione Markdown → struttura dati...")
    converter = MarkdownToDocxConverter()
    elements = converter.parse_markdown(markdown_content)

    print(f"📄 Generazione documento DOCX ({len(elements)} elementi)...")
    doc = converter.convert_to_docx(elements)

    doc.save(str(output_path))
    print(f"✅ Documento creato: {output_path}")

    return output_path


def main():
    """Entry point"""
    try:
        if not SOURCE_MD.exists():
            raise FileNotFoundError(f"File Markdown non trovato: {SOURCE_MD}")

        result = convert_markdown_to_docx(SOURCE_MD)
        print(f"\n🎉 Conversione completata!")
        print(f"   Input: {SOURCE_MD.name}")
        print(f"   Output: {result.name}")
        print(f"   Metodo: Generazione da zero senza template")

    except Exception as e:
        print(f"\n❌ Errore: {e}")
        return 1

    return 0


if __name__ == "__main__":
    exit(main())
